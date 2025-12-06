"""
Word文档导出工具类
用于将研究计划的完整分析过程导出为Word文档
"""

import os
from datetime import datetime

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


class PlanningWordExporter:
    """研究计划Word文档导出器"""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """设置文档样式"""
        # 标题样式
        title_style = self.doc.styles.add_style("CustomTitle", WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = "微软雅黑"
        title_font.size = Pt(16)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)

        # 二级标题样式
        heading_style = self.doc.styles.add_style(
            "CustomHeading", WD_STYLE_TYPE.PARAGRAPH
        )
        heading_font = heading_style.font
        heading_font.name = "微软雅黑"
        heading_font.size = Pt(14)
        heading_font.bold = True
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)

        # 正文样式
        body_style = self.doc.styles.add_style("CustomBody", WD_STYLE_TYPE.PARAGRAPH)
        body_font = body_style.font
        body_font.name = "宋体"
        body_font.size = Pt(11)
        body_style.paragraph_format.line_spacing = 1.2
        body_style.paragraph_format.space_after = Pt(6)

    def export_planning_results(
        self,
        original_question,
        optimized_queries,
        context,
        individual_plans,
        integrated_plan,
        output_path=None,
    ):
        """
        导出完整的研究计划分析结果

        Args:
            original_question: 原始研究问题
            optimized_queries: 优化后的查询列表
            context: RAG检索的上下文结果
            individual_plans: 各个优化问题对应的计划列表
            integrated_plan: 整合后的最终计划
            output_path: 输出文件路径，如果为None则自动生成
        """
        # 预处理所有输入数据，清理XML不兼容字符
        original_question = self._clean_xml_content(
            str(original_question) if original_question else ""
        )
        context = self._clean_xml_content(str(context) if context else "")
        integrated_plan = self._clean_xml_content(
            str(integrated_plan) if integrated_plan else ""
        )

        # 清理列表数据
        if optimized_queries:
            optimized_queries = [
                self._clean_xml_content(str(q)) for q in optimized_queries
            ]
        else:
            optimized_queries = []

        if individual_plans:
            individual_plans = [
                self._clean_xml_content(str(p)) for p in individual_plans
            ]
        else:
            individual_plans = []

        # 文档标题
        title = self.doc.add_paragraph("研究计划分析报告", style="CustomTitle")

        # 生成时间
        time_info = self.doc.add_paragraph(
            f"生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}"
        )
        time_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()  # 空行

        # 1. 原始研究问题
        self._add_section_title("1. 原始研究问题")
        self._add_content(original_question)

        # 2. 优化后的查询需求
        self._add_section_title("2. 优化后的查询需求")
        for i, query in enumerate(optimized_queries, 1):
            self._add_content(f"查询 {i}: {query}")

        # 3. 相关领域知识
        self._add_section_title("3. 相关领域知识与上下文")
        self._add_content(context)

        # 4. 各查询对应的研究计划
        self._add_section_title("4. 各查询对应的研究计划")
        for i, plan in enumerate(individual_plans, 1):
            query_title = (
                optimized_queries[i - 1]
                if i - 1 < len(optimized_queries)
                else f"计划 {i}"
            )
            self._add_subsection_title(f"4.{i} 基于查询: {query_title[:50]}...")
            self._add_content(plan)

        # 5. 整合后的最终计划
        self._add_section_title("5. 整合后的最终研究计划")
        self._add_content(integrated_plan)

        # 保存文档
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"research_planning_report_{timestamp}.docx"

        self.doc.save(output_path)
        return output_path

    def _add_section_title(self, title):
        """添加章节标题"""
        clean_title = self._clean_xml_content(title)
        self.doc.add_paragraph(clean_title, style="CustomHeading")

    def _add_subsection_title(self, title):
        """添加小节标题"""
        clean_title = self._clean_xml_content(title)
        p = self.doc.add_paragraph()
        run = p.add_run(clean_title)
        run.font.name = "微软雅黑"
        run.font.size = Pt(12)
        run.font.bold = True
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

    def _add_content(self, content, max_length=None):
        """添加正文内容"""
        # 清理字符串，移除NULL字节和控制字符
        content = self._clean_xml_content(content)

        if max_length and len(content) > max_length:
            content = content[:max_length] + "\n\n...(内容过长，已截取前{}字符)".format(
                max_length
            )

        # 处理长文本，按段落分割
        paragraphs = content.split("\n\n")
        for para in paragraphs:
            if para.strip():
                # 再次清理每个段落
                clean_para = self._clean_xml_content(para.strip())
                self.doc.add_paragraph(clean_para, style="CustomBody")

        self.doc.add_paragraph()  # 添加空行分隔

    def _clean_xml_content(self, text):
        """清理文本内容，移除XML不兼容的字符"""
        if not isinstance(text, str):
            text = str(text)

        # 移除NULL字节
        text = text.replace("\x00", "")

        # 移除其他控制字符（保留常用的换行符、制表符等）
        import re

        # 保留常用的空白字符：\t(制表符), \n(换行), \r(回车), 空格
        # 移除其他控制字符（ASCII 0-31，除了9,10,13）
        text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", "", text)

        return text


def export_planning_to_word(
    original_question,
    optimized_queries,
    context,
    individual_plans,
    integrated_plan,
    output_dir="output",
):
    """
    便捷函数：导出研究计划到Word文档

    Args:
        original_question: 原始问题
        optimized_queries: 优化查询列表
        context: 上下文
        individual_plans: 个别计划列表
        integrated_plan: 整合计划
        output_dir: 输出目录

    Returns:
        str: 生成的文件路径
    """
    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)

    # 生成文件路径
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"research_planning_report_{timestamp}.docx"
    output_path = os.path.join(output_dir, filename)

    # 创建导出器并导出
    exporter = PlanningWordExporter()
    return exporter.export_planning_results(
        original_question=original_question,
        optimized_queries=optimized_queries,
        context=context,
        individual_plans=individual_plans,
        integrated_plan=integrated_plan,
        output_path=output_path,
    )


if __name__ == "__main__":
    """测试Word文档导出功能"""
    print("=== Word文档导出工具测试 ===")

    # 模拟测试数据
    test_original_question = "Predict broadly neutralizing antibodies against H5N1 influenza virus using single-cell V(D)J data, identify conserved B cell subsets, and reveal structural features of neutralization breadth"

    test_optimized_queries = [
        "Integrate single-cell V(D)J sequencing data using Scanpy to identify conserved B cell subsets, then predict broadly neutralizing antibodies against H5N1 influenza using the MetaBCR model trained on influenza viruses",
        "Apply Harmony for batch effect correction in single-cell data integration, followed by differential gene expression analysis with Seurat to uncover B cell subset markers and correlate with antibody neutralization breadth using the MetaBCR model",
        "Utilize AlphaFold3 to predict the 3D structure of H5N1 viral proteins and assess receptor-ligand interactions with MetaBCR, while performing pseudotime inference on single-cell data using Scanpy to trace B cell differentiation and neutralization potential",
    ]

    test_context = """ResouRce NATURE ImmUNoLoGy (Extended Data Fig. 1c; Methods) were corrected for technical 10× run batch effects using BBKNN 20 or Harmony 21. Based on the BBKNN-corrected data, unsupervised clustering followed by a two-dimensional uniform manifold approximation and projection (UMAP)22 revealed 20 molecularly distinct clusters (Fig. 1b, 27 clusters before batch correction, and Extended Data Fig. 1d,e). Similar results were obtained using Harmony21, which yielded 21 clusters (Extended Data Fig. 1f). Cell numbers in clusters varied from 42,353 to 624 (Extended Data Fig. 2a). Clusters were defined on the basis of gene expression values compared to all other cells (Fig. 1c and Supplementary Table 2a). This analysis identified single clusters of dendritic cells (DCs), megakaryocytes (Mgk), erythroid cells (Eryth), natural killer (NK) cells, plasma blasts/cells (PBs/PCs) and plasmacytoid DCs (pDCs), two clusters of B cells and monocytes (CD14+ and CD16 + ), five CD4+ T cell clusters and five CD8+ T cell clusters."""

    test_individual_plans = [
        """### 基于Scanpy的H5N1广谱中和抗体预测计划

#### Step 1: 数据预处理
- 使用Scanpy加载单细胞V(D)J数据
- 质量控制：过滤低质量细胞和基因
- 标准化和对数变换

#### Step 2: B细胞亚群识别
- 使用leiden聚类算法识别B细胞亚群
- 标记基因分析识别保守亚群
- UMAP可视化B细胞分布

#### Step 3: MetaBCR H5N1模型预测
- 提取配对的重链和轻链序列
- 应用Meta-BCR Influenza模型进行中和预测
- 筛选高置信度的广谱中和抗体候选

#### Step 4: 结果验证
- 交叉验证预测结果
- 与已知H5N1中和抗体比较""",
        """### 基于Harmony批次校正的综合分析计划

#### Step 1: 批次效应校正
- 使用Harmony校正不同样本间的批次效应
- 评估校正效果和数据质量

#### Step 2: Seurat差异基因分析
- 使用Seurat进行B细胞亚群标记基因识别
- 差异表达分析找出关键调控基因
- 功能富集分析

#### Step 3: MetaBCR整合分析
- 结合转录组数据和BCR序列信息
- 使用Meta-BCR模型预测中和能力
- 关联基因表达与中和广度

#### Step 4: 生物学验证
- 设计实验验证关键发现
- ELISA和中和实验确认预测结果""",
        """### 基于AlphaFold3的结构功能分析计划

#### Step 1: 结构预测
- 使用AlphaFold3预测H5N1 HA蛋白结构
- 预测候选抗体的3D结构
- 评估结构质量和可信度

#### Step 2: 分子对接分析
- 进行抗体-抗原分子对接
- 分析结合界面和关键残基
- 计算结合亲和力

#### Step 3: 轨迹推断分析
- 使用Scanpy进行B细胞发育轨迹分析
- 追踪中和抗体产生的分化路径
- 识别关键发育节点

#### Step 4: 整合分析
- 结合结构、功能和发育数据
- 构建综合的中和机制模型""",
    ]

    test_integrated_plan = """### H5N1广谱中和抗体预测与B细胞亚群分析整合计划

#### Step 1: 数据预处理与质量控制
- 使用Scanpy加载和预处理单细胞V(D)J数据
- 应用Harmony进行批次效应校正
- 质量控制：过滤低质量细胞，保留高质量BCR序列

#### Step 2: B细胞亚群鉴定与特征分析
- 使用Leiden聚类算法识别B细胞亚群
- Seurat差异基因表达分析识别亚群标记基因
- UMAP可视化B细胞分布和发育轨迹

#### Step 3: 保守B细胞亚群筛选
- 跨样本分析识别保守的B细胞克隆型
- 功能富集分析确定亚群生物学功能
- 筛选具有中和潜力的B细胞亚群

#### Step 4: MetaBCR H5N1模型预测
- 提取配对的重链和轻链序列
- 应用Meta-BCR Influenza模型进行广谱中和预测
- 结合转录组数据提高预测准确性

#### Step 5: 结构功能验证
- 使用AlphaFold3预测候选抗体和H5N1 HA蛋白结构
- 分子对接分析抗体-抗原相互作用
- 识别关键的中和表位和结构特征

#### Step 6: 实验验证与优化
- 设计实验验证预测的广谱中和抗体
- ELISA和假病毒中和实验确认活性
- 基于结果优化预测模型"""

    try:
        # 测试导出功能
        output_path = export_planning_to_word(
            original_question=test_original_question,
            optimized_queries=test_optimized_queries,
            context=test_context,
            individual_plans=test_individual_plans,
            integrated_plan=test_integrated_plan,
            output_dir="test_output",
        )

        print(f"✅ 测试成功！")
        print(f"📄 Word文档已生成: {output_path}")
        print(f"📁 文件大小: {os.path.getsize(output_path) / 1024:.1f} KB")
        print(f"🕐 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # 验证文件是否存在
        if os.path.exists(output_path):
            print(f"✅ 文件验证通过，可以正常打开")
        else:
            print(f"❌ 文件生成失败")

    except ImportError as e:
        print(f"❌ 依赖包缺失: {e}")
        print("请安装依赖: pip install python-docx")
    except Exception as e:
        print(f"❌ 测试失败: {e}")
        import traceback

        traceback.print_exc()

    print("\n=== 测试完成 ===")
